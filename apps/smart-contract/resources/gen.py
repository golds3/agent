from docx import Document

doc = Document()
doc.add_heading("合约规则说明文档", level=1)

# 劳动合同
doc.add_heading("一、劳动合同", level=2)
doc.add_paragraph("1. 合同期限：分为固定期限、无固定期限和完成一定工作任务为期限。")
doc.add_paragraph("2. 工作内容与地点：明确员工职责与工作地点。")
doc.add_paragraph("3. 工作时间与休息：标准为每日不超过8小时，每周不超过40小时。")
doc.add_paragraph("4. 劳动报酬：工资发放方式与构成。")
doc.add_paragraph("5. 社保：依法缴纳五险一金。")

# 服务外包合同
doc.add_heading("二、服务外包合同", level=2)
doc.add_paragraph("1. 合同主体：发包方与承包方信息。")
doc.add_paragraph("2. 服务范围：服务内容与标准。")
doc.add_paragraph("3. 服务费用：金额、支付方式与时间。")
doc.add_paragraph("4. 保密条款：对商业信息保密。")
doc.add_paragraph("5. 验收标准：服务验收后付款。")

# 采购合同
doc.add_heading("三、采购合同", level=2)
doc.add_paragraph("1. 标的与数量：产品名称、规格、数量、交货时间。")
doc.add_paragraph("2. 质量要求：质量或检测标准。")
doc.add_paragraph("3. 价格与付款：单价、总额、付款方式。")
doc.add_paragraph("4. 交货与验收：地点、方式、方法与时限。")
doc.add_paragraph("5. 违约责任：逾期、质量不合格等违约条款。")

doc.save("合约规则说明文档.docx")
